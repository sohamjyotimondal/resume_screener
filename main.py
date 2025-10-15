"""
Main module for resume processing - handles file reading and text extraction
Provides class-based API for integration with Flask/FastAPI backends
"""

import PyPDF2
import docx
from typing import List, Dict, Optional, Union, BinaryIO
import logging
from pathlib import Path
import tempfile
import os
from parser import ResumeParser
from screener import ResumeScreener

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeExtractor:
    """Handles extraction of text and URLs from resume files (PDF, DOCX)."""

    @staticmethod
    def extract_urls_from_pdf(file_path: str) -> List[str]:
        """Extract URLs from PDF hyperlinks/annotations."""
        urls = []
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    # Check if page has annotations
                    if "/Annots" in page:
                        annotations = page["/Annots"]
                        for annotation in annotations:
                            obj = annotation.get_object()
                            # Check for Link annotations with URLs
                            if obj.get("/Subtype") == "/Link":
                                if "/A" in obj:  # Action
                                    action = obj["/A"]
                                    if "/URI" in action:
                                        url = action["/URI"]
                                        if url and url not in urls:
                                            urls.append(str(url))
        except Exception as e:
            logger.warning(f"Could not extract URLs from PDF: {e}")

        return urls

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text and URLs from PDF file."""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            # Also extract URLs from hyperlinks
            urls = ResumeExtractor.extract_urls_from_pdf(file_path)
            if urls:
                text += "\n\nEXTRACTED URLS/LINKS:\n"
                for url in urls:
                    text += f"- {url}\n"

            return text

    @staticmethod
    def extract_urls_from_docx(file_path: str) -> List[str]:
        """Extract URLs from DOCX hyperlinks."""
        urls = []
        try:
            doc = docx.Document(file_path)
            # Get all hyperlinks from the document relationships
            rels = doc.part.rels
            for rel in rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel.target_ref
                    if url and url not in urls:
                        # Filter out internal anchors (starting with #)
                        if not url.startswith("#"):
                            urls.append(url)
        except Exception as e:
            logger.warning(f"Could not extract URLs from DOCX: {e}")

        return urls

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text and URLs from DOCX file."""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Also extract URLs from hyperlinks
        urls = ResumeExtractor.extract_urls_from_docx(file_path)
        if urls:
            text += "\n\nEXTRACTED URLS/LINKS:\n"
            for url in urls:
                text += f"- {url}\n"

        return text

    @staticmethod
    def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
        """Extract text and URLs from PDF bytes (for file uploads)."""
        import io

        file_obj = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"

        # Note: URL extraction from bytes is limited, works best with file path
        return text

    @staticmethod
    def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
        """Extract text and URLs from DOCX bytes (for file uploads)."""
        import io

        file_obj = io.BytesIO(file_bytes)
        doc = docx.Document(file_obj)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text

    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """
        Extract text from a resume file (PDF or DOCX).
        Auto-detects file type based on extension.
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path_obj.suffix.lower()

        if extension == ".pdf":
            logger.info(f"Extracting text from PDF: {file_path}")
            return ResumeExtractor.extract_text_from_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            logger.info(f"Extracting text from DOCX: {file_path}")
            return ResumeExtractor.extract_text_from_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. Supported formats: .pdf, .docx, .doc"
            )

    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        """
        Extract text from file bytes (for file uploads).
        Auto-detects file type based on filename extension.
        """
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            logger.info(f"Extracting text from uploaded PDF: {filename}")
            return ResumeExtractor.extract_text_from_pdf_bytes(file_bytes)
        elif extension in [".docx", ".doc"]:
            logger.info(f"Extracting text from uploaded DOCX: {filename}")
            return ResumeExtractor.extract_text_from_docx_bytes(file_bytes)
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. Supported formats: .pdf, .docx, .doc"
            )


class ResumeProcessor:
    """
    Main class for processing resumes - parsing and screening.
    Designed for easy integration with Flask/FastAPI backends.
    """

    def __init__(
        self,
        parser_model: str = "llama-3.3-70b-versatile",
        screener_model: str = "llama-3.3-70b-versatile",
    ):
        """
        Initialize the resume processor.

        Args:
            parser_model: Groq model for parsing resumes
            screener_model: Groq model for screening resumes
        """
        self.parser = ResumeParser(model=parser_model)
        self.screener = ResumeScreener(model=screener_model)
        self.extractor = ResumeExtractor()

    def parse_resume_from_path(self, file_path: str) -> Dict:
        """
        Parse resume from file path.

        Args:
            file_path: Path to resume file

        Returns:
            Parsed resume as dictionary
        """
        resume_text = self.extractor.extract_text_from_file(file_path)
        logger.info(f"Extracted {len(resume_text)} characters from {file_path}")

        parsed_resume = self.parser.parse_resume(resume_text)
        logger.info("Resume parsed successfully")

        return parsed_resume.model_dump(exclude_none=True)

    def parse_resume_from_bytes(self, file_bytes: bytes, filename: str) -> Dict:
        """
        Parse resume from file bytes (for file uploads).

        Args:
            file_bytes: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Parsed resume as dictionary
        """
        resume_text = self.extractor.extract_text_from_bytes(file_bytes, filename)
        logger.info(
            f"Extracted {len(resume_text)} characters from uploaded file: {filename}"
        )

        parsed_resume = self.parser.parse_resume(resume_text)
        logger.info("Resume parsed successfully")

        return parsed_resume.model_dump(exclude_none=True)

    def screen_resume(
        self,
        parsed_resume: Dict,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Screen a parsed resume against job requirements.

        Args:
            parsed_resume: Parsed resume dictionary
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring categories

        Returns:
            Screening result as dictionary
        """
        screening_result = self.screener.screen_resume(
            parsed_resume, job_title, job_description, weights
        )
        logger.info(
            f"Resume screened. Overall score: {screening_result.overall_score}/10"
        )

        return screening_result.model_dump(exclude_none=True)

    def process_resume_from_path(
        self,
        file_path: str,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Complete workflow: Parse and screen resume from file path.

        Args:
            file_path: Path to resume file
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring

        Returns:
            Dictionary with 'parsed' and 'screened' keys
        """
        logger.info(f"Processing resume from path: {file_path}")

        # Parse resume
        parsed = self.parse_resume_from_path(file_path)

        # Screen resume
        screened = self.screen_resume(parsed, job_title, job_description, weights)

        return {"parsed": parsed, "screened": screened}

    def process_resume_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        job_title: str,
        job_description: str,
        weights: Optional[Dict[str, float]] = None,
    ) -> Dict:
        """
        Complete workflow: Parse and screen resume from file bytes (for uploads).
        Perfect for Flask/FastAPI endpoints.

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            job_title: Job position title
            job_description: Job description and requirements
            weights: Optional custom weights for scoring

        Returns:
            Dictionary with 'parsed' and 'screened' keys
        """
        logger.info(f"Processing uploaded resume: {filename}")

        # Parse resume
        parsed = self.parse_resume_from_bytes(file_bytes, filename)

        # Screen resume
        screened = self.screen_resume(parsed, job_title, job_description, weights)

        return {"parsed": parsed, "screened": screened}


if __name__ == "__main__":
    import sys

    # Initialize processor
    processor = ResumeProcessor()

    # Check if user wants to run screening workflow or just parsing
    if len(sys.argv) > 1 and sys.argv[1] == "screen":
        # Complete screening workflow
        print("=" * 60)
        print("Resume Screening - Complete Workflow")
        print("=" * 60)

        try:
            resume_file = "resume_f.pdf"
            job_title = "Senior Software Engineer"
            job_description = """
            We are looking for a Senior Software Engineer with:
            - 5+ years of software development experience
            - Strong proficiency in Python, JavaScript, or Java
            - Experience with cloud platforms (AWS, GCP, or Azure)
            - Knowledge of microservices architecture
            - Experience with CI/CD pipelines
            - Strong problem-solving and communication skills
            - Bachelor's degree in Computer Science or related field
            """

            print(f"\n📄 Resume: {resume_file}")
            print(f"💼 Position: {job_title}")

            # Process resume (parse + screen)
            result = processor.process_resume_from_path(
                resume_file, job_title, job_description
            )

            # Display screening results
            screened = result["screened"]
            print("\n" + "=" * 60)
            print("SCREENING RESULTS")
            print("=" * 60)
            print(f"\n⭐ Overall Score: {screened['overall_score']:.1f}/10")
            print(f"📋 Recommendation: {screened['recommendation']}")
            print(f"\n💬 Summary:\n{screened['summary']}")

            print(f"\n{'='*60}")
            print("CATEGORY SCORES")
            print(f"{'='*60}")
            print(f"🛠️  Skills: {screened['skill_match']['score']:.1f}/10")
            print(f"💼 Experience: {screened['experience_match']['score']:.1f}/10")
            print(f"🎓 Education: {screened['education_match']['score']:.1f}/10")
            print(f"🚀 Projects: {screened['project_match']['score']:.1f}/10")
            print(
                f"📜 Certifications: {screened['certification_match']['score']:.1f}/10"
            )
            print(f"🤝 Cultural Fit: {screened['cultural_fit']['score']:.1f}/10")

            print(f"\n{'='*60}")
            print("✅ STRENGTHS")
            print(f"{'='*60}")
            for i, strength in enumerate(screened["strengths"], 1):
                print(f"{i}. {strength}")

            print(f"\n{'='*60}")
            print("⚠️  CONCERNS")
            print(f"{'='*60}")
            for i, concern in enumerate(screened["concerns"], 1):
                print(f"{i}. {concern}")

            print(f"\n{'='*60}")
            print(f"📌 Next Steps: {screened['next_steps']}")
            print("=" * 60)

        except Exception as e:
            print(f"\n❌ Error: {e}")
            import traceback

            traceback.print_exc()

    else:
        # Just parse the resume
        print("=" * 60)
        print("Resume Parser - Complete Workflow")
        print("=" * 60)

        try:
            resume_file = "resume_f.pdf"
            print(f"\n📄 Processing: {resume_file}")

            parsed_data = processor.parse_resume_from_path(resume_file)

            # Save to JSON
            import json

            output_file = "parsed_resume.json"
            with open(output_file, "w") as f:
                json.dump(parsed_data, f, indent=2)

            print(f"\n✅ Results saved to: {output_file}")
            print("=" * 60)

        except FileNotFoundError:
            print(f"\n❌ Error: File not found")
        except Exception as e:
            print(f"\n❌ Error: {e}")
            import traceback

            traceback.print_exc()
